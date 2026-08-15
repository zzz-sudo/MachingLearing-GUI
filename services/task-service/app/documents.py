from __future__ import annotations

from pathlib import Path

import pymupdf
from docx import Document
from openpyxl import Workbook
from rapidocr import RapidOCR

from app.errors import document_export_error, import_error
from app.models import DocumentPage, DocumentParseResult
from app.storage import WorkspaceStore


MIN_TEXT_LAYER_LENGTH = 20
OCR_SCALE = 2

EXPORT_MEDIA_TYPES = {
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "md": "text/markdown; charset=utf-8",
    "txt": "text/plain; charset=utf-8",
}


class DocumentExportService:
    def __init__(self, store: WorkspaceStore) -> None:
        self.store = store

    def export(self, asset_id: str, output_format: str) -> tuple[Path, str]:
        asset, _source_path = self.store.resolve_asset_path(asset_id)
        project = self.store.get_project(asset.project_id)
        result = self.store.get_document_result(asset_id)
        output_directory = Path(project.path) / "reports" / "document-exports"
        output_directory.mkdir(parents=True, exist_ok=True)
        output_path = output_directory / f"{Path(asset.name).stem}-parsed.{output_format}"

        try:
            if output_format == "docx":
                self._write_docx(output_path, asset.name, result)
            elif output_format == "xlsx":
                self._write_xlsx(output_path, asset.name, result)
            elif output_format == "md":
                output_path.write_text(self._as_markdown(result), encoding="utf-8")
            elif output_format == "txt":
                output_path.write_text(self._as_text(result), encoding="utf-8")
            else:
                raise ValueError(f"不支持的导出格式: {output_format}")
        except (OSError, ValueError) as error:
            raise document_export_error(asset_id, output_format, str(error)) from error

        return output_path, EXPORT_MEDIA_TYPES[output_format]

    @staticmethod
    def _as_markdown(result: DocumentParseResult) -> str:
        return "\n\n".join(
            f"## 第 {page.page_number} 页\n\n{page.text}"
            for page in result.pages
        )

    @staticmethod
    def _as_text(result: DocumentParseResult) -> str:
        return "\n\n".join(
            f"第 {page.page_number} 页\n{page.text}"
            for page in result.pages
        )

    @staticmethod
    def _write_docx(path: Path, source_name: str, result: DocumentParseResult) -> None:
        document = Document()
        document.add_heading(Path(source_name).stem, level=0)
        document.add_paragraph(f"来源文件: {source_name}")
        document.add_paragraph(f"解析引擎: {result.engine}")
        for page in result.pages:
            document.add_heading(f"第 {page.page_number} 页", level=1)
            document.add_paragraph(page.text or "本页没有可导出的文本")
        document.save(path)

    @staticmethod
    def _write_xlsx(path: Path, source_name: str, result: DocumentParseResult) -> None:
        workbook = Workbook()
        summary = workbook.active
        summary.title = "文档信息"
        summary.append(["来源文件", source_name])
        summary.append(["解析引擎", result.engine])
        summary.append(["文档类型", result.pdf_type])
        summary.append(["页数", result.page_count])
        pages = workbook.create_sheet("页面内容")
        pages.append(["页码", "文本", "需要OCR"])
        for page in result.pages:
            pages.append([page.page_number, page.text, page.needs_ocr])
        pages.column_dimensions["A"].width = 10
        pages.column_dimensions["B"].width = 100
        pages.column_dimensions["C"].width = 14
        workbook.save(path)


class PdfDocumentService:
    def __init__(self) -> None:
        self._ocr: RapidOCR | None = None

    def parse(self, path: Path, asset_id: str, project_root: Path) -> DocumentParseResult:
        try:
            document = pymupdf.open(path)
        except (pymupdf.FileDataError, OSError) as error:
            raise import_error(
                "DocumentParseError",
                f"无法打开 PDF: {path.name}",
                "pdf_parse",
                filename=path.name,
                reason=str(error),
            ) from error

        pages: list[DocumentPage] = []
        ocr_pages: list[int] = []
        try:
            for index, page in enumerate(document):
                text = page.get_text("text").strip()
                page_number = index + 1
                needs_ocr = len(text) < MIN_TEXT_LAYER_LENGTH
                if needs_ocr:
                    ocr_pages.append(page_number)
                    text = self._extract_ocr_text(page)
                    needs_ocr = not bool(text)
                pages.append(
                    DocumentPage(
                        page_number=page_number,
                        text=text,
                        needs_ocr=needs_ocr,
                    )
                )
        except Exception as error:
            raise import_error(
                "DocumentParseError",
                f"PDF 页面解析失败: {path.name}",
                "pdf_parse",
                filename=path.name,
                reason=str(error),
            ) from error
        finally:
            document.close()

        pages_needing_ocr = [page.page_number for page in pages if page.needs_ocr]
        if len(ocr_pages) == len(pages):
            pdf_type = "scanned"
        elif ocr_pages:
            pdf_type = "mixed"
        else:
            pdf_type = "text_based"

        if not pages_needing_ocr:
            status = "parsed"
        elif len(pages_needing_ocr) == len(pages):
            status = "ocr_required"
        else:
            status = "partial"

        markdown = "\n\n".join(
            f"## 第 {page.page_number} 页\n\n{page.text}"
            for page in pages
            if page.text
        )
        documents_dir = project_root / "documents"
        base_name = f"{path.stem}-{asset_id.removeprefix('asset-')[:8]}"
        markdown_path = documents_dir / f"{base_name}.md"
        json_path = documents_dir / f"{base_name}.json"
        engine = "pymupdf+rapidocr" if ocr_pages else "pymupdf"
        result = DocumentParseResult(
            asset_id=asset_id,
            pdf_type=pdf_type,
            engine=engine,
            status=status,
            page_count=len(pages),
            ocr_pages=ocr_pages,
            pages_needing_ocr=pages_needing_ocr,
            markdown_relative_path=(
                markdown_path.relative_to(project_root).as_posix() if markdown else None
            ),
            json_relative_path=json_path.relative_to(project_root).as_posix(),
            markdown_preview=markdown[:4000],
            pages=pages,
        )
        if markdown:
            markdown_path.write_text(markdown, encoding="utf-8")
        json_path.write_text(
            result.model_dump_json(by_alias=True, indent=2),
            encoding="utf-8",
        )
        return result

    def _extract_ocr_text(self, page: pymupdf.Page) -> str:
        if self._ocr is None:
            self._ocr = RapidOCR()
        pixmap = page.get_pixmap(
            matrix=pymupdf.Matrix(OCR_SCALE, OCR_SCALE),
            alpha=False,
        )
        result = self._ocr(pixmap.tobytes("png"))
        return "\n".join(result.txts or ()).strip()
